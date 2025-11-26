import os
from typing import Optional
import PyPDF2
from docx import Document

# pdfplumber is optional (requires Rust compilation, may fail on some platforms)
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

def extract_text_from_file(file_path: str, file_ext: str) -> Optional[str]:
    """
    Extract text from various file formats
    """
    try:
        if file_ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return extract_text_from_docx(file_path)
        elif file_ext == '.doc':
            # Note: .doc files require additional libraries like python-docx2txt or antiword
            # For now, we'll try to use docx converter or suggest conversion
            return extract_text_from_doc(file_path)
        elif file_ext == '.txt':
            return extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")
    except Exception as e:
        raise Exception(f"Error extracting text from {file_ext} file: {str(e)}")

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using pdfplumber (if available) or PyPDF2
    """
    text = ""
    
    # Try pdfplumber first if available (better for complex layouts)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception:
            pass
    
    # Use PyPDF2 (always available)
    if not text.strip():
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_doc(file_path: str) -> str:
    """
    Extract text from DOC file
    Note: This is a basic implementation. For better results, consider using:
    - python-docx2txt
    - antiword (requires system installation)
    - or convert .doc to .docx first
    """
    # For now, raise an error suggesting conversion
    # In production, you might want to use docx2txt or other libraries
    raise Exception(
        ".doc files are not fully supported. Please convert to .docx or .pdf. "
        "You can install python-docx2txt for better .doc support."
    )

def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from TXT file
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    except Exception as e:
        raise Exception(f"Failed to extract text from TXT: {str(e)}")




